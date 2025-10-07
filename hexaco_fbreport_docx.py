# -*- coding: utf-8 -*-
import os
import re
import io
from datetime import datetime, timezone, timedelta
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from openai import OpenAI
client = OpenAI(api_key="")

# ------------------ コンフィグ ------------------
CSV_PATH = "csv/20250417_nttdata_ddd.csv"
TEMPLATE_PERSON = "tmp/HEXACOfbレポート_本人用_tmp.docx"
TEMPLATE_OFFICE = "tmp/HEXACOfbレポート_事務局用_tmp.docx"
OUT_DIR = "out/"
os.makedirs(OUT_DIR, exist_ok=True)

# 画像サイズ（高さ px）— ここを変えるだけで出力サイズを統一変更できます
RADAR_HEIGHT_PX = 300

# フォント（テンプレ内テキストの標準フォントとして適用を試みます）
FONT_NAME = "MS Gothic"

# ------------------ ユーティリティ ------------------

def jst_today_str():
    jst = timezone(timedelta(hours=9))
    return datetime.now(jst).strftime("%Y/%m/%d")

def sanitize_filename(name: str) -> str:
    safe = re.sub(r'[\\/*?:\"<>|]+', "_", str(name))
    safe = safe.strip()
    return safe or "unknown"

def label_from_score(x: float) -> str:
    if pd.isna(x):
        return "中"
    if x >= 4.0:
        return "高い"
    elif x < 2.5:
        return "低い"
    else:
        return "中"

def trim_to_fullwidth_chars(text: str, limit: int) -> str:
    if text is None:
        return ""
    s = text.strip()
    if len(s) <= limit:
        return s
    s = s[:limit]
    last_marume = max(s.rfind("。"), s.rfind("！"), s.rfind("？"))
    if last_marume >= 0:
        s = s[:last_marume+1]
    return s

def detect_strength_weakness(row: pd.Series, start_col: int = 7):
    strengths, weaknesses = [], []
    for col in row.index[start_col:]:
        val = str(row[col]).strip().lower()
        if val == "high":
            strengths.append(col)
        elif val == "low":
            weaknesses.append(col)
    return strengths, weaknesses

# ------------------ レーダー（メモリ返却） ------------------

def make_radar_chart_buffer(values, labels, height_px=None):
    """PNGをBytesIOで返す（ファイルは保存しない）"""
    if height_px is None:
        height_px = RADAR_HEIGHT_PX

    N = len(values)
    angles = np.linspace(0, 2*np.pi, N, endpoint=False).tolist()
    values = list(values) + [values[0]]
    angles += [angles[0]]

    fig = plt.figure(figsize=(4, 4))
    ax = fig.add_subplot(111, polar=True)
    ax.set_theta_offset(np.pi / 2)
    ax.set_theta_direction(-1)
    ax.set_thetagrids(np.degrees(angles[:-1]), labels)
    ax.set_rlabel_position(0)
    ax.set_ylim(0, 5)
    ax.plot(angles, values, linewidth=2)
    ax.fill(angles, values, alpha=0.25)
    plt.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf

# ------------------ テンプレ置換 ------------------

def replace_text_placeholders(doc: Document, mapping: dict):
    def _replace_in_paragraph(paragraph, mapping):
        if not paragraph.text:
            return
        full_text = paragraph.text
        replaced = False
        for key, val in mapping.items():
            if key in full_text:
                full_text = full_text.replace(key, val)
                replaced = True
        if replaced:
            for run in paragraph.runs[::-1]:
                paragraph._p.remove(run._r)
            paragraph.add_run(full_text)

    for p in doc.paragraphs:
        _replace_in_paragraph(p, mapping)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_in_paragraph(p, mapping)

def replace_image_placeholder(doc: Document, placeholder: str, image_source, height_px=None):
    """image_source は BytesIO でもファイルパス文字列でも可"""
    if height_px is None:
        height_px = RADAR_HEIGHT_PX

    targets = []
    for p in doc.paragraphs:
        if placeholder in p.text:
            targets.append(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if placeholder in p.text:
                        targets.append(p)
    for p in targets:
        for run in p.runs[::-1]:
            p._p.remove(run._r)
        height_inch = height_px / 96.0
        run = p.add_run()
        run.add_picture(image_source, height=Inches(height_inch))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ------------------ フォント適用 ------------------

def apply_font(doc: Document, font_name: str):
    """文書の既定スタイルと段落・表セルのランにフォント名を適用（EastAsiaも設定）"""
    try:
        # 既定（Normal）スタイル
        style = doc.styles["Normal"]
        style.font.name = font_name
        # East Asia用フォント指定
        if style._element.rPr is None:
            style._element._new_rPr()
        rFonts = style._element.rPr.rFonts
        rFonts.set(qn("w:eastAsia"), font_name)
    except Exception:
        pass  # テンプレ側にスタイルがなければ黙って無視

    # 既存の全ランに適用（明示的に上書きする）
    def set_run_font(run):
        try:
            run.font.name = font_name
            r = run._r
            r.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        except Exception:
            pass

    for p in doc.paragraphs:
        for run in p.runs:
            set_run_font(run)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        set_run_font(run)

# ------------------ コメント生成 ------------------

PERSON_COMMENT_LIMIT = 250
OFFICE_COMMENT_LIMIT  = 600 

def build_person_prompt(name: str, levels_5: dict, strengths: list, weaknesses: list) -> str:
    """
    本人用（O, C, E, A, Nの5因子）コメント生成プロンプトを返す。
    levels_5 のキー: O, C, E, A, N （値は「高い/中/低」）
    strengths/weaknesses は CSV の high/low 列名（最大10件まで使用）
    """
    lines = []
    lines.append("あなたは産業・組織心理学の専門家です。全角250文字以内で日本語の自然な文を書いてください。")
    lines.append(f"対象者名: {name}")
    lines.append("HEXACO（本人用5因子）の水準: " + ", ".join([
        f"O={levels_5.get('O','')}",
        f"C={levels_5.get('C','')}",
        f"E={levels_5.get('E','')}",
        f"A={levels_5.get('A','')}",
        f"N={levels_5.get('N','')}",
    ]))
    if strengths:
        lines.append("強み（high）: " + "、".join(strengths[:10]))
    if weaknesses:
        lines.append("改善余地（low）: " + "、".join(weaknesses[:10]))
    # 本人向けのトーンと要件
    lines.append("要件: 否定表現を避け、前向きな行動提案を3個か4個ぐらい含め、専門用語を控えめに、1段落のみ。")
    return "\n".join(lines)

def build_office_prompt(name: str, levels_6: dict, strengths: list, weaknesses: list) -> str:
    """
    事務局用（H, E, X, A, C, Oの6因子）コメント生成プロンプトを返す。
    levels_6 のキー: H, E, X, A, C, O （値は「高い/中/低」）
    strengths/weaknesses は CSV の high/low 列名（最大10件まで使用）
    """
    lines = []
    lines.append("あなたは産業・組織心理学の専門家です。全角600文字以内で日本語の自然な文を書いてください。")
    lines.append(f"対象者名: {name}")
    lines.append("HEXACO（事務局用6因子）の水準: " + ", ".join([
        f"H={levels_6.get('H','')}",
        f"E={levels_6.get('E','')}",
        f"X={levels_6.get('X','')}",
        f"A={levels_6.get('A','')}",
        f"C={levels_6.get('C','')}",
        f"O={levels_6.get('O','')}",
    ]))
    if strengths:
        lines.append("強み（high）: " + "、".join(strengths[:10]))
    if weaknesses:
        lines.append("改善余地（low）: " + "、".join(weaknesses[:10]))
    # 事務局向けのトーンと要件
    lines.append("要件: 人事・教育担当者向けに配置・育成上の示唆を含め、客観的かつ簡潔に、専門用語を控えめに、1段落でお願いします。特にダークトライアドの傾向がある場合は注意を強調してください。")
    return "\n".join(lines)

def generate_comment_via_gpt(prompt: str) -> str:
    try:
        resp = client.responses.create(
            # model="gpt-5",
            model="gpt-4o-mini",
            input=prompt,
            temperature=0.7,
            max_output_tokens=512,
        )
        return resp.output_text.strip()
    except Exception:
        return "観察された特性を踏まえ、強みを活かしつつ小さな行動から改善を進めましょう。"

# ------------------ DOCX生成 ------------------

def fill_person_docx(row: pd.Series, radar_buf, out_path: str):
    """本人用（Hなし／O,C,E,A,Nの5因子）"""
    doc = Document(TEMPLATE_PERSON)

    name = str(row.get("Name", "NoName"))
    levels = {
        "O": label_from_score(row.get("開放性（好奇心）")),
        "C": label_from_score(row.get("誠実性（計画性）")),
        "E": label_from_score(row.get("外向性（ポジティブさ）")),
        "A": label_from_score(row.get("協調性（利他性・共感性）")),
        "N": label_from_score(row.get("情動性（不安傾向）")),
    }
    strengths, weaknesses = detect_strength_weakness(row, start_col=7)

    text_map = {
        "[Name]": name,
        "[YYYY/MM/DD]": jst_today_str(),
        "[reputate_hexaco_O]": levels["O"],
        "[reputate_hexaco_C]": levels["C"],
        "[reputate_hexaco_E]": levels["E"],
        "[reputate_hexaco_A]": levels["A"],
        "[reputate_hexaco_N]": levels["N"],
        "[LEVEL_O]": levels["O"],
        "[LEVEL_C]": levels["C"],
        "[LEVEL_E]": levels["E"],
        "[LEVEL_A]": levels["A"],
        "[LEVEL_N]": levels["N"],
    }
    replace_text_placeholders(doc, text_map)

    prompt = build_person_prompt(name, levels, strengths, weaknesses)
    comment = generate_comment_via_gpt(prompt)
    comment = trim_to_fullwidth_chars(comment, PERSON_COMMENT_LIMIT)
    replace_text_placeholders(doc, {"[comment_about_5_factors]": comment, "[COMMENT]": comment})

    for key in ["[radar_chart_5_factors_height200px]", "[RADAR_5]", "[radar_chart]"]:
        radar_buf.seek(0)
        replace_image_placeholder(doc, key, radar_buf, height_px=RADAR_HEIGHT_PX)

    # フォント適用
    apply_font(doc, FONT_NAME)

    doc.save(out_path)

def fill_office_docx(row: pd.Series, radar_buf, out_path: str):
    """事務局用（6因子）"""
    doc = Document(TEMPLATE_OFFICE)

    name = str(row.get("Name", "NoName"))
    levels = {
        "H": label_from_score(row.get("正直・謙虚さ（倫理観）")),
        "E": label_from_score(row.get("情動性（不安傾向）")),
        "X": label_from_score(row.get("外向性（ポジティブさ）")),
        "A": label_from_score(row.get("協調性（利他性・共感性）")),
        "C": label_from_score(row.get("誠実性（計画性）")),
        "O": label_from_score(row.get("開放性（好奇心）")),
    }
    strengths, weaknesses = detect_strength_weakness(row, start_col=7)

    text_map = {
        "[Name]": name,
        "[YYYY/MM/DD]": jst_today_str(),
        "[reputate_hexaco_H]": levels["H"],
        "[reputate_hexaco_E]": levels["E"],
        "[reputate_hexaco_X]": levels["X"],
        "[reputate_hexaco_A]": levels["A"],
        "[reputate_hexaco_C]": levels["C"],
        "[reputate_hexaco_O]": levels["O"],
        "[LEVEL_H]": levels["H"],
        "[LEVEL_E]": levels["E"],
        "[LEVEL_X]": levels["X"],
        "[LEVEL_A]": levels["A"],
        "[LEVEL_C]": levels["C"],
        "[LEVEL_O]": levels["O"],
    }
    replace_text_placeholders(doc, text_map)

    prompt = build_office_prompt(name, levels, strengths, weaknesses)
    comment = generate_comment_via_gpt(prompt)
    comment = trim_to_fullwidth_chars(comment, OFFICE_COMMENT_LIMIT)
    replace_text_placeholders(doc, {"[comment_about_6_factors_and_darktrait]": comment, "[COMMENT]": comment})

    for key in ["[radar_chart_6_factors_height200px]", "[RADAR_6]", "[radar_chart]"]:
        radar_buf.seek(0)
        replace_image_placeholder(doc, key, radar_buf, height_px=RADAR_HEIGHT_PX)

    # フォント適用
    apply_font(doc, FONT_NAME)

    doc.save(out_path)

# ------------------ メイン ------------------

def main():
    df = pd.read_csv(CSV_PATH, na_values=["NA", "N/A", "na", "NaN", "-", ""], encoding="utf-8")

    # 本人用（5因子）: O, C, E, A, N
    person_cols = ["開放性（好奇心）", "誠実性（計画性）", "外向性（ポジティブさ）", "協調性（利他性・共感性）", "情動性（不安傾向）"]
    # 事務局用（6因子）: H, E, X, A, C, O
    office_cols  = ["正直・謙虚さ（倫理観）", "情動性（不安傾向）", "外向性（ポジティブさ）", "協調性（利他性・共感性）", "誠実性（計画性）", "開放性（好奇心）"]

    for c in set(person_cols + office_cols):
        if c in df.columns:
            df[c] = pd.to_numeric(df[c], errors="coerce").clip(lower=0, upper=5)

    for idx, row in df.iterrows():
        name = str(row.get("Name", f"row{idx+1}"))
        safe_name = sanitize_filename(name)

        # レーダー（本人用・5因子）
        vals_p = [row.get(c, np.nan) for c in person_cols]
        s_p = pd.Series(vals_p, dtype="float64")
        filled_p = s_p.fillna(s_p.mean() if not np.isnan(s_p.mean()) else 0).tolist()
        buf_p = make_radar_chart_buffer(filled_p, ["O", "C", "E", "A", "N"], height_px=RADAR_HEIGHT_PX)

        # レーダー（事務局用・6因子）
        vals_o = [row.get(c, np.nan) for c in office_cols]
        s_o = pd.Series(vals_o, dtype="float64")
        filled_o = s_o.fillna(s_o.mean() if not np.isnan(s_o.mean()) else 0).tolist()
        buf_o = make_radar_chart_buffer(filled_o, ["H", "E", "X", "A", "C", "O"], height_px=RADAR_HEIGHT_PX)

        # DOCX生成
        fill_person_docx(row, buf_p, os.path.join(OUT_DIR, f"{safe_name}_本人用.docx"))
        fill_office_docx(row, buf_o, os.path.join(OUT_DIR, f"{safe_name}_事務局用.docx"))
        print(f"Generated: {os.path.join(OUT_DIR, f'{safe_name}_本人用.docx')}")
        print(f"Generated: {os.path.join(OUT_DIR, f'{safe_name}_事務局用.docx')}")

if __name__ == "__main__":
    main()
