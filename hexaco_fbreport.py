# -*- coding: utf-8 -*-
import os
import re
import io
from datetime import datetime, timezone, timedelta
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from typing import Dict, List, Tuple, Iterable, Optional
import sys
from pathlib import Path
import json

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx2pdf import convert

try:
    from openai import OpenAI
except Exception as e:
    print("`openai` パッケージの読み込みに失敗しました。`pip install openai` を実行してください。", file=sys.stderr)
    raise

openai_client = OpenAI()

# -------- OpenAI Model & Prompt Files (added) --------
# MODEL = "gpt-5-nano"
MODEL = "gpt-4.1"
# MODEL = "gpt-4o-mini"
# MODEL = "o4-mini"
PROMPT_PERSONAL_FILE = "pmt/prompt_personal.txt"
PROMPT_OFFICE_FILE = "pmt/prompt_office.txt"

# ------------------ コンフィグ ------------------
# CSV_PATH = "csv/20250417_nttdata_ddd.csv"
# CSV_PATH = "csv/20251020_nttdatauniv_test1.csv"
# CSV_PATH = "csv/20251020_nttdatauniv_test2.csv"
CSV_PATH = "csv/20251020_nttdatauniv.csv"
TEMPLATE_PERSON = "tmp/HEXACOfbレポート_本人用_tmp.docx"
TEMPLATE_OFFICE = "tmp/HEXACOfbレポート_事務局用_tmp.docx"
OUT_DIR = "out/"
OUT_PERSON_WORD = os.path.join(OUT_DIR, "本人用/word")
OUT_PERSON_PDF  = os.path.join(OUT_DIR, "本人用/pdf")
OUT_OFFICE_WORD = os.path.join(OUT_DIR, "事務局用/word")
OUT_OFFICE_PDF  = os.path.join(OUT_DIR, "事務局用/pdf")
for d in [OUT_PERSON_WORD, OUT_PERSON_PDF, OUT_OFFICE_WORD, OUT_OFFICE_PDF]:
    os.makedirs(d, exist_ok=True)

# 画像サイズ（高さ px）— ここを変えるだけで出力サイズを統一変更できます
RADAR_HEIGHT_PX = 300

# フォント（テンプレ内テキストの標準フォントとして適用を試みます）
FONT_NAME = "MS Gothic"

DARK_TRAIT_COLS = ["ダーク傾向", "ナルシシズム", "サイコパシー", "マキャベリズム"]

# ------------------ ユーティリティ ------------------

def jst_today_str():
    jst = timezone(timedelta(hours=9))
    return datetime.now(jst).strftime("%Y/%m/%d")

def sanitize_filename(name: str) -> str:
    safe = re.sub(r'[\\/*?:\"<>|]+', "_", str(name))
    safe = safe.strip()
    return safe or "unknown"

def label_from_score(x: float) -> str:
    if pd.isna(x):
        return "中"
    if x >= 3.8:
        return "高い"
    elif x < 2.5:
        return "低い"
    else:
        return "中"

def fmt1(x) -> str:
    """小数第1位の文字列（元データの小数第1位）"""
    try:
        return f"{float(x):.1f}"
    except Exception:
        return ""

def trim_to_fullwidth_chars(text: str, limit: int) -> str:
    if text is None:
        return ""
    s = text.strip()
    if len(s) <= limit:
        return s
    s = s[:limit]
    last_marume = max(s.rfind("。"), s.rfind("！"), s.rfind("？"))
    if last_marume >= 0:
        s = s[:last_marume+1]
    return s

def collect_all_level_flags(
    row: pd.Series,
    include_values: tuple[str, ...] = ("high", "middle", "low"),
    exclude_cols: list[str] | set[str] | None = None,
) -> list[str]:
    """
    CSVの行から、値が high/middle/low の列を「カラム:値」形式で収集。
    例: ["疲れやすさ:low", "主体性:high", ...]
    exclude_cols に指定された列名は除外（DARK_TRAIT_COLSなど）。
    """
    out = []
    iv = {s.lower() for s in include_values}
    ex = set(exclude_cols or [])

    for col, val in row.items():
        if col in ex:
            continue
        if isinstance(val, str):
            sv = val.strip().lower()
            if sv in iv:
                out.append(f"{col}:{sv}")
    return out

# ------------------ レーダー ------------------

def make_radar_chart_buffer(values, labels, height_px=None, overlay=None, overlay_style=None, fill_alpha=0.25):
    """
    PNGをBytesIOで返す（ファイルは保存しない）
    - values: 主系列（個人）
    - overlay: オーバーレイ系列（例：全体平均）。None可
    - overlay_style: dict（例：{'color':'red','linewidth':2}）
    """
    if height_px is None:
        height_px = RADAR_HEIGHT_PX

    N = len(values)
    angles = np.linspace(0, 2*np.pi, N, endpoint=False).tolist()

    vals_main = list(values) + [values[0]]
    angles2 = angles + [angles[0]]

    fig = plt.figure(figsize=(4, 4))
    ax = fig.add_subplot(111, polar=True)
    ax.set_theta_offset(np.pi / 2)
    ax.set_theta_direction(-1)
    ax.set_thetagrids(np.degrees(angles), labels)
    ax.set_rlabel_position(0)
    ax.set_ylim(0, 5)

    # main
    ax.plot(angles2, vals_main, linewidth=2)
    ax.fill(angles2, vals_main, alpha=fill_alpha)

    # overlay
    if overlay is not None:
        vals_overlay = list(overlay) + [overlay[0]]
        style = overlay_style or {}
        ax.plot(angles2, vals_overlay, **style)

    plt.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf

# ------------------ テンプレ置換 ------------------

def replace_text_placeholders(doc: Document, mapping: dict):
    def _replace_in_paragraph(paragraph, mapping):
        if not paragraph.text:
            return
        full_text = paragraph.text
        replaced = False
        for key, val in mapping.items():
            if key in full_text:
                full_text = full_text.replace(key, val)
                replaced = True
        if replaced:
            for run in paragraph.runs[::-1]:
                paragraph._p.remove(run._r)
            paragraph.add_run(full_text)

    for p in doc.paragraphs:
        _replace_in_paragraph(p, mapping)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_in_paragraph(p, mapping)

def replace_image_placeholder(doc: Document, placeholder: str, image_source, height_px=None):
    """image_source は BytesIO でもファイルパス文字列でも可"""
    if height_px is None:
        height_px = RADAR_HEIGHT_PX

    targets = []
    for p in doc.paragraphs:
        if placeholder in p.text:
            targets.append(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if placeholder in p.text:
                        targets.append(p)
    for p in targets:
        for run in p.runs[::-1]:
            p._p.remove(run._r)
        height_inch = height_px / 96.0
        run = p.add_run()
        run.add_picture(image_source, height=Inches(height_inch))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ------------------ フォント適用 ------------------

def apply_font(doc: Document, font_name: str):
    """文書の既定スタイルと段落・表セルのランにフォント名を適用（EastAsiaも設定）"""
    try:
        style = doc.styles["Normal"]
        style.font.name = font_name
        if style._element.rPr is None:
            style._element._new_rPr()
        rFonts = style._element.rPr.rFonts
        rFonts.set(qn("w:eastAsia"), font_name)
    except Exception:
        pass

    def set_run_font(run):
        try:
            run.font.name = font_name
            r = run._r
            r.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        except Exception:
            pass

    for p in doc.paragraphs:
        for run in p.runs:
            set_run_font(run)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        set_run_font(run)

# ------------------ コメント生成（APIは空） ------------------

PERSON_COMMENT_LIMIT = 200
OFFICE_COMMENT_LIMIT  = 450

MAX_OFFICE_STRENGTHS = 10
MAX_OFFICE_WEAKNESSES = 10


def build_person_prompt(name: str, scores: dict, levels: dict) -> str:
    """
    name: 受検者名（例: "山田太郎"）
    scores: {"O": 4.3, "C": 3.5, "E": 2.6, "A": 3.7, "N": 3.1}
    levels: {"O": "high"|"middle"|"low", ... for O,C,E,A,N}
    """

    prompt_path = Path(__file__).resolve().parent / PROMPT_PERSONAL_FILE
    if not prompt_path.exists():
        print(f"本人用の固定プロンプトファイルが見つかりません: {prompt_path}", file=sys.stderr)
        return "テンプレートが見つかりません"

    template = prompt_path.read_text(encoding="utf-8")

    # JSON文字列（日本語保持・コンパクト化）
    payload = {"name": name, "scores": scores, "levels": levels}
    payload_json = json.dumps(payload, ensure_ascii=False, separators=(",", ":"))
    scores_json  = json.dumps(scores, ensure_ascii=False, separators=(",", ":"))
    levels_json  = json.dumps(levels, ensure_ascii=False, separators=(",", ":"))

    # format() では他プレースホルダ干渉の恐れがあるため、ピンポイント置換
    prompt = template.replace("{payload}", payload_json)\
                     .replace("{name}", name)\
                     .replace("{scores}", scores_json)\
                     .replace("{levels}", levels_json)
    return prompt.strip()


def build_office_prompt(
    name: str,
    levels_6: dict,
    level_flags: list[str],
    dark_levels: dict | None = None
) -> str:

    payload = {
        "name": name,
        "hexaco_levels": levels_6,
        "level_flags": level_flags,
        "dark_trait_levels": dark_levels or {},
    }

    # ← 文字列化（テンプレがプレーン埋め込み想定なのでJSON文字列を使う）
    levels_6_str     = json.dumps(levels_6, ensure_ascii=False, separators=(",", ":"))
    level_flags_str  = json.dumps(level_flags, ensure_ascii=False, separators=(",", ":"))
    dark_levels_str  = json.dumps(dark_levels or {}, ensure_ascii=False, separators=(",", ":"))

    prompt_path = Path(__file__).resolve().parent / PROMPT_OFFICE_FILE
    if not prompt_path.exists():
        print(f"事務局用の固定プロンプトファイルが見つかりません: {prompt_path}", file=sys.stderr)
        return "テンプレートが見つかりません"

    template = prompt_path.read_text(encoding="utf-8")

    prompt = (
        template
        .replace("{name}", name)
        .replace("{hexaco_levels}", levels_6_str)
        .replace("{level_flags}", level_flags_str)
        .replace("{dark_trait_levels}", dark_levels_str)
    )
    return prompt

def generate_comment_via_gpt(prompt: str) -> str:
    try:
        resp = openai_client.responses.create(
            model=MODEL,
            input=[
                {"role": "system", "content": "あなたは簡潔かつ丁寧な日本語の文章アシスタントです。"},
                {"role": "user", "content": prompt},
            ]
        )
        return resp.output_text.strip()
    except KeyboardInterrupt:
        print("\nユーザーにより中断されました。", file=sys.stderr)
        return 130
    except Exception as e:
        # 予期せぬエラー
        print(f"[ERROR] {type(e).__name__}: {e}", file=sys.stderr)
        return "観察された特性を踏まえ、強みを活かしつつ小さな行動から改善を進めましょう。"


# ------------------ DOCX生成 ------------------

def fill_person_docx(row: pd.Series, radar_buf, out_docx_path: str, out_pdf_path: str):
    """本人用（Hなし／O,C,E,A,Nの5因子）"""
    doc = Document(TEMPLATE_PERSON)

    name = str(row.get("Name", "NoName"))
    raw_vals = {
        "O": fmt1(row.get("開放性（好奇心）")),
        "C": fmt1(row.get("勤勉性（計画性）")),
        "E": fmt1(row.get("外向性（ポジティブさ）")),
        "A": fmt1(row.get("協調性（利他性・共感性）")),
        "N": fmt1(row.get("情動性（不安傾向）")),
    }
    levels = {
        "O": label_from_score(row.get("開放性（好奇心）")),
        "C": label_from_score(row.get("勤勉性（計画性）")),
        "E": label_from_score(row.get("外向性（ポジティブさ）")),
        "A": label_from_score(row.get("協調性（利他性・共感性）")),
        "N": label_from_score(row.get("情動性（不安傾向）")),
    }
    text_map = {
        "[Name]": name,
        "[YYYY/MM/DD]": jst_today_str(),
        "[reputate_hexaco_O]": levels["O"],
        "[reputate_hexaco_C]": levels["C"],
        "[reputate_hexaco_E]": levels["E"],
        "[reputate_hexaco_A]": levels["A"],
        "[reputate_hexaco_N]": levels["N"],
        "[LEVEL_O]": levels["O"],
        "[LEVEL_C]": levels["C"],
        "[LEVEL_E]": levels["E"],
        "[LEVEL_A]": levels["A"],
        "[LEVEL_N]": levels["N"],
        "[value_hexaco_O]": raw_vals["O"],
        "[value_hexaco_C]": raw_vals["C"],
        "[value_hexaco_E]": raw_vals["E"],
        "[value_hexaco_A]": raw_vals["A"],
        "[value_hexaco_N]": raw_vals["N"],
    }
    replace_text_placeholders(doc, text_map)

    def to_float(v):
        if v is None or v == "":
            return None
        try:
            return float(v)
        except Exception:
            # 余裕があればここでログ出し
            return None

    scores = {k: to_float(v) for k, v in raw_vals.items()}

    name = str(row.get("Name", "")).strip()

    # コメント
    prompt = build_person_prompt(name=name, scores=scores, levels=levels)

    comment = generate_comment_via_gpt(prompt)
    comment = trim_to_fullwidth_chars(comment, PERSON_COMMENT_LIMIT)
    replace_text_placeholders(doc, {"[comment_about_5_factors]": comment, "[COMMENT]": comment})

    # レーダー画像
    for key in ["[radar_chart_5_factors_height200px]", "[RADAR_5]", "[radar_chart]"]:
        radar_buf.seek(0)
        replace_image_placeholder(doc, key, radar_buf, height_px=RADAR_HEIGHT_PX)

    # フォント適用
    apply_font(doc, FONT_NAME)

    doc.save(out_docx_path)
    convert(out_docx_path, out_pdf_path)

def fill_office_docx(row: pd.Series, radar_buf, out_docx_path: str, out_pdf_path: str):
    """事務局用（6因子）"""
    doc = Document(TEMPLATE_OFFICE)

    name = str(row.get("Name", "NoName"))
    # 元データの素点（1桁小数）を直接埋め込む
    raw_vals = {
        "H": fmt1(row.get("正直・謙虚さ（倫理観）")),
        "E": fmt1(row.get("情動性（不安傾向）")),
        "X": fmt1(row.get("外向性（ポジティブさ）")),
        "A": fmt1(row.get("協調性（利他性・共感性）")),
        "C": fmt1(row.get("勤勉性（計画性）")),
        "O": fmt1(row.get("開放性（好奇心）")),
    }
    levels = {
        "H": label_from_score(row.get("正直・謙虚さ（倫理観）")),
        "E": label_from_score(row.get("情動性（不安傾向）")),
        "X": label_from_score(row.get("外向性（ポジティブさ）")),
        "A": label_from_score(row.get("協調性（利他性・共感性）")),
        "C": label_from_score(row.get("勤勉性（計画性）")),
        "O": label_from_score(row.get("開放性（好奇心）")),
    }

    level_flags = collect_all_level_flags(row, exclude_cols=DARK_TRAIT_COLS)

    text_map = {
        "[Name]": name,
        "[YYYY/MM/DD]": jst_today_str(),
        "[reputate_hexaco_H]": levels["H"],
        "[reputate_hexaco_E]": levels["E"],
        "[reputate_hexaco_X]": levels["X"],
        "[reputate_hexaco_A]": levels["A"],
        "[reputate_hexaco_C]": levels["C"],
        "[reputate_hexaco_O]": levels["O"],
        "[LEVEL_H]": levels["H"],
        "[LEVEL_E]": levels["E"],
        "[LEVEL_X]": levels["X"],
        "[LEVEL_A]": levels["A"],
        "[LEVEL_C]": levels["C"],
        "[LEVEL_O]": levels["O"],
        "[value_hexaco_H]": raw_vals["H"],
        "[value_hexaco_E]": raw_vals["E"],
        "[value_hexaco_X]": raw_vals["X"],
        "[value_hexaco_A]": raw_vals["A"],
        "[value_hexaco_C]": raw_vals["C"],
        "[value_hexaco_O]": raw_vals["O"],
    }
    replace_text_placeholders(doc, text_map)

    dark_levels = {}
    for col in DARK_TRAIT_COLS:
        if col in row.index:
            dark_levels[col] = row[col]

    # コメント
    prompt = build_office_prompt(name, levels_6=levels, level_flags=level_flags, dark_levels=dark_levels)
    comment = generate_comment_via_gpt(prompt)
    comment = trim_to_fullwidth_chars(comment, OFFICE_COMMENT_LIMIT)
    replace_text_placeholders(doc, {"[comment_about_6_factors_and_darktrait]": comment, "[COMMENT]": comment})

    # レーダー画像（事務局用は main 内で平均オーバーレイ済みのバッファを受け取る）
    for key in ["[radar_chart_6_factors_height200px]", "[RADAR_6]", "[radar_chart]"]:
        radar_buf.seek(0)
        replace_image_placeholder(doc, key, radar_buf, height_px=RADAR_HEIGHT_PX)

    # フォント適用
    apply_font(doc, FONT_NAME)

    doc.save(out_docx_path)
    convert(out_docx_path, out_pdf_path)

# ------------------ メイン ------------------

def main():
    df = pd.read_csv(CSV_PATH, na_values=["NA", "N/A", "na", "NaN", "-", ""], encoding="utf-8")

    DROP_COLS_EXACT = [
        "価値観の傾向：人や資源を管理し、お金を求める",
        "価値観の傾向：社会的に認められた成功を求める",
        "価値観の傾向：快楽を求める",
        "価値観の傾向：刺激的な経験を求める",
        "価値観の傾向：思考と行動の独立性を求める",
        "価値観の傾向：平等や社会的正義や環境保護を求める",
        "価値観の傾向：周りの人々の繁栄や幸福を求める",
        "価値観の傾向：他人の期待に応えるために自らの衝動をコントロールする",
        "価値観の傾向：伝統を守る",
        "価値観の傾向：自分・家族・国家の安全や安心を求める",
        "開放性と正直謙虚さが高い人と相性がいい可能性",
        "開放性が高く正直謙虚さが低い人と相性がいい可能性",
        "開放性が低く正直謙虚さが高い人と相性がいい可能性",
        "開放性と正直謙虚さが低い人と相性がいい可能性"
    ]
    exact_hits = [c for c in DROP_COLS_EXACT if c in df.columns]

    cols_to_drop = sorted(set(exact_hits))
    if cols_to_drop:
        df = df.drop(columns=cols_to_drop, errors="ignore")

    # 本人用（5因子）: O, C, E, A, N
    person_cols = ["開放性（好奇心）", "勤勉性（計画性）", "外向性（ポジティブさ）", "協調性（利他性・共感性）", "情動性（不安傾向）"]
    # 事務局用（6因子）: H, E, X, A, C, O
    office_cols  = ["正直・謙虚さ（倫理観）", "情動性（不安傾向）", "外向性（ポジティブさ）", "協調性（利他性・共感性）", "勤勉性（計画性）", "開放性（好奇心）"]

    # 数値化＆クリップ
    for c in set(person_cols + office_cols):
        if c in df.columns:
            df[c] = pd.to_numeric(df[c], errors="coerce").clip(lower=0, upper=5)

    # 事務局用の全体平均（6因子の順で）
    avg_series = df[office_cols].mean(numeric_only=True)
    avg_vals = avg_series.fillna(avg_series.mean() if not np.isnan(avg_series.mean()) else 0).tolist()

    for idx, row in df.iterrows():
        name = str(row.get("Name", f"row{idx+1}"))
        safe_name = sanitize_filename(name)

        # レーダー（本人用・5因子）
        vals_p = [row.get(c, np.nan) for c in person_cols]
        s_p = pd.Series(vals_p, dtype="float64")
        filled_p = s_p.fillna(s_p.mean() if not np.isnan(s_p.mean()) else 0).tolist()
        buf_p = make_radar_chart_buffer(
            filled_p,
            ["O", "C", "E", "A", "N"],
            height_px=RADAR_HEIGHT_PX
        )

        # レーダー（事務局用・6因子） — 全体平均を赤線でオーバーレイ
        vals_o = [row.get(c, np.nan) for c in office_cols]
        s_o = pd.Series(vals_o, dtype="float64")
        filled_o = s_o.fillna(s_o.mean() if not np.isnan(s_o.mean()) else 0).tolist()
        buf_o = make_radar_chart_buffer(
            filled_o,
            ["H", "E", "X", "A", "C", "O"],
            height_px=RADAR_HEIGHT_PX,
            overlay=avg_vals,
            overlay_style={"color": "red", "linewidth": 2},
            fill_alpha=0.20,
        )

        # ---- パスを作る（4フォルダに振り分ける）----
        person_docx = os.path.join(OUT_PERSON_WORD, f"{safe_name}_本人用.docx")
        person_pdf  = os.path.join(OUT_PERSON_PDF,  f"{safe_name}_本人用.pdf")
        office_docx = os.path.join(OUT_OFFICE_WORD, f"{safe_name}_事務局用.docx")
        office_pdf  = os.path.join(OUT_OFFICE_PDF,  f"{safe_name}_事務局用.pdf")

        # ---- 出力 ----
        # fill_person_docx(row, buf_p, person_docx, person_pdf)
        fill_office_docx(row, buf_o, office_docx, office_pdf)

        print(f"Generated: {person_docx}")
        print(f"Generated: {person_pdf}")
        print(f"Generated: {office_docx}")
        print(f"Generated: {office_pdf}")

if __name__ == "__main__":
    main()
